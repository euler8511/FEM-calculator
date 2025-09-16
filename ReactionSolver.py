import numpy as np
import meshio
import pyvista as pv
from scipy.sparse import lil_matrix
from scipy.sparse.linalg import spsolve
import os
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class ForceAnalysis:
    """
    Gmsh 메쉬 파일을 읽어 3D 탄성학 유한요소해석을 수행하고, 결과를 시각화하며,
    최종 리포트를 생성하는 클래스.
    - 10절점 사면체(Tetra10) 요소를 사용합니다.
    - 희소 행렬(Sparse Matrix)을 사용하여 메모리 효율성을 높였습니다.
    - GUI로부터 전달받은 개별 점(Point) 하중 및 고정 조건을 적용합니다.
    """
    def __init__(self, msh_file, force_data, fix_data, E, v):
        """
        클래스 초기화
        :param msh_file: Gmsh 메쉬 파일 경로 (.msh)
        :param force_data: 힘의 크기와 위치 정보가 담긴 딕셔너리
        :param fix_data: 고정할 위치와 자유도 정보가 담긴 리스트
        :param E: 탄성 계수 (Young's Modulus)
        :param v: 푸아송 비 (Poisson's Ratio)
        """
        self.msh_file = msh_file
        self.force_data = force_data
        self.fix_data = fix_data
        self.E = E
        self.v = v
        self.pd = 3  # 3차원 문제
        
        # 해석에 필요한 변수들 초기화
        self.points = None
        self.tetra10_conn = None
        self.num_nodes = 0
        self.K = None
        self.u = None
        self.f = None
        self.reaction_forces = None
        self.fixed_nodes_info = [] # 고정점 정보 저장을 위한 리스트
        self.negative_detJ_count = 0 # [NEW] 음수 자코비안 카운터
        
        # 메쉬 파일 읽기
        self._read_mesh()
        # 재료 강성 행렬 생성
        self._create_material_matrix()

    def _read_mesh(self):
        """meshio를 사용하여 메쉬 파일을 읽고 주요 정보를 멤버 변수에 저장합니다."""
        print("1. Reading mesh file...")
        self.mesh = meshio.read(self.msh_file)
        self.points = self.mesh.points
        self.num_nodes = len(self.points)
        
        # 3D 체적 요소 (Tetra10) 추출
        self.tetra10_conn = self.mesh.cells_dict.get('tetra10')
        if self.tetra10_conn is None:
            raise ValueError("메쉬 파일에 'tetra10' 요소가 없습니다.")
        
        # Physical Group에서 노드 인덱스 추출
        self.diri_nodes = self._get_nodes_from_physical_group('Diri_BCs', 'vertex')
        self.neumann_nodes = self._get_nodes_from_physical_group('Neumann_BCs', 'vertex')

        print(f"   - Nodes: {self.num_nodes}, Tetra10 Elements: {len(self.tetra10_conn)}")

    def _get_nodes_from_physical_group(self, group_name, cell_type):
        """지정된 Physical Group에 속한 모든 노드의 인덱스를 반환합니다."""
        try:
            group_id = self.mesh.field_data[group_name][0]
            phys_tags = self.mesh.cell_data_dict['gmsh:physical'][cell_type]
            group_cell_indices = np.where(phys_tags == group_id)[0]
            group_cells = self.mesh.cells_dict[cell_type][group_cell_indices]
            return np.unique(group_cells.flatten())
        except (KeyError, IndexError):
            print(f"Warning: Physical group '{group_name}' or cell type '{cell_type}' not found in mesh.")
            return np.array([], dtype=int)

    def _create_material_matrix(self):
        """등방성 선형 탄성 재료의 강성 행렬(C)을 생성합니다."""
        C1 = self.E / ((1 + self.v) * (1 - 2 * self.v))
        C2 = (1 - 2 * self.v) / 2
        self.C = C1 * np.array([
            [1-self.v, self.v,     self.v,     0,   0,   0],
            [self.v,     1-self.v, self.v,     0,   0,   0],
            [self.v,     self.v,     1-self.v, 0,   0,   0],
            [0,       0,       0,       C2,  0,   0],
            [0,       0,       0,       0,   C2,  0],
            [0,       0,       0,       0,   0,   C2]
        ])

    def _shape_funcs_tet10(self, xi, eta, zeta):
        """Tetra10 요소의 형상 함수 및 자연 좌표계 기준 도함수를 반환합니다."""
        L2, L3, L4 = xi, eta, zeta
        L1 = 1 - xi - eta - zeta
        dN_L = np.array([
            [4*L1-1, 0,    0,    0   ], [0, 4*L2-1,    0,    0   ],
            [0,    0, 4*L3-1,    0   ], [0,    0,    0, 4*L4-1 ],
            [4*L2, 4*L1,    0,    0   ], [0, 4*L3, 4*L2,    0   ],
            [4*L3, 0, 4*L1,    0   ], [4*L4, 0,    0, 4*L1 ],
            [0, 4*L4, 0, 4*L2 ], [0,    0, 4*L4, 4*L3 ]
        ]).T
        dL = np.array([[-1,-1,-1], [1,0,0], [0,1,0], [0,0,1]])
        dN = dL.T @ dN_L
        return dN

    def assemble_stiffness_matrix(self):
        """전체 강성 행렬(K)을 조립합니다."""
        print("2. Assembling global stiffness matrix (K)...")
        total_dof = self.pd * self.num_nodes
        self.K = lil_matrix((total_dof, total_dof))
        gauss_pts = np.array([[0.58541020, 0.13819660, 0.13819660],
                              [0.13819660, 0.58541020, 0.13819660],
                              [0.13819660, 0.13819660, 0.58541020],
                              [0.13819660, 0.13819660, 0.13819660]])
        w = 1/4

        for tet_nodes in self.tetra10_conn:
            el_coords = self.points[tet_nodes]
            Ke = np.zeros((30, 30))
            for pt in gauss_pts:
                dN_natural = self._shape_funcs_tet10(pt[0], pt[1], pt[2])
                J = dN_natural @ el_coords
                detJ = np.linalg.det(J)
                if detJ <= 1e-12:
                    self.negative_detJ_count += 1 # [MODIFIED] 음수 자코비안 카운트
                    continue
                
                dN_global = np.linalg.inv(J) @ dN_natural
                B = np.zeros((6, 30))
                for i in range(10):
                    dx, dy, dz = dN_global[:, i]
                    idx = 3 * i
                    B[0, idx] = dx; B[1, idx+1] = dy; B[2, idx+2] = dz
                    B[3, idx] = dy; B[3, idx+1] = dx
                    B[4, idx+1] = dz; B[4, idx+2] = dy
                    B[5, idx] = dz; B[5, idx+2] = dx
                Ke += B.T @ self.C @ B * detJ * w

            dof_indices = (self.pd * np.repeat(tet_nodes, self.pd) + np.tile(range(self.pd), 10)).flatten()
            self.K[np.ix_(dof_indices, dof_indices)] += Ke
        
        self.K = self.K.tocsr()
        print("   - Assembly complete.")

    def apply_boundary_conditions(self):
        """GUI에서 받은 점(Point) 기반 경계 조건을 적용합니다."""
        print("3. Applying point-based boundary conditions...")
        total_dof = self.pd * self.num_nodes
        self.f = np.zeros(total_dof)
        fixed_dofs = []
        self.fixed_nodes_info = []

        for fix_info in self.fix_data:
            pos = np.array([fix_info['pos_x'], fix_info['pos_y'], fix_info['pos_z']])
            distances = np.linalg.norm(self.points[self.diri_nodes] - pos, axis=1)
            node_idx = self.diri_nodes[np.argmin(distances)]
            dofs = []
            if fix_info['fix_x'] == 0: dofs.append(3 * node_idx)
            if fix_info['fix_y'] == 0: dofs.append(3 * node_idx + 1)
            if fix_info['fix_z'] == 0: dofs.append(3 * node_idx + 2)
            fixed_dofs.extend(dofs)
            self.fixed_nodes_info.append({'node_idx': node_idx, 'pos': self.points[node_idx], 'dofs': dofs})

        self.fixed_dofs = np.unique(fixed_dofs)
        print(f"   - Fixed {len(self.fixed_dofs)} DOFs.")

        force_vec = [self.force_data['force_x'], self.force_data['force_y'], self.force_data['force_z']]
        pos = np.array([self.force_data['force_x_pstn'], self.force_data['force_y_pstn'], self.force_data['force_z_pstn']])
        distances = np.linalg.norm(self.points[self.neumann_nodes] - pos, axis=1)
        node_idx = self.neumann_nodes[np.argmin(distances)]
        self.f[3 * node_idx : 3 * node_idx + 3] = force_vec
        print(f"   - Applied force {force_vec} N to node {node_idx}.")
        
        self.active_dofs = np.setdiff1d(np.arange(total_dof), self.fixed_dofs)

    def solve(self):
        """선형 시스템을 풀어 변위를 계산하고 반력을 계산합니다."""
        print("4. Solving the linear system...")
        K_reduced = self.K[np.ix_(self.active_dofs, self.active_dofs)]
        f_reduced = self.f[self.active_dofs]
        u_reduced = spsolve(K_reduced, f_reduced)
        self.u = np.zeros(self.pd * self.num_nodes)
        self.u[self.active_dofs] = u_reduced
        print("   - System solved.")
        self.reaction_forces = self.K @ self.u
        
    def print_reactions(self):
        """계산된 반력을 콘솔에 출력합니다."""
        if self.reaction_forces is None: return
        print("\n--- Reaction Forces ---")
        total_reaction = np.zeros(3)
        for i, info in enumerate(self.fixed_nodes_info):
            node_idx = info['node_idx']
            reactions = self.reaction_forces[3 * node_idx : 3 * node_idx + 3]
            total_reaction += reactions
            print(f"  Node {node_idx} (Fix Point {i+1}): Rx={reactions[0]:.4e}, Ry={reactions[1]:.4e}, Rz={reactions[2]:.4e} N")
        print("\n--- Force Equilibrium Check ---")
        applied_force = np.array([self.force_data['force_x'], self.force_data['force_y'], self.force_data['force_z']])
        print(f"  Sum of Applied Forces (Fx, Fy, Fz): {applied_force}")
        print(f"  Sum of Reaction Forces (Rx, Ry, Rz): {-total_reaction}")

    def run_simulation(self):
        """전체 해석 파이프라인을 실행하고 리포트를 생성합니다."""
        self.assemble_stiffness_matrix()
        self.apply_boundary_conditions()
        self.solve()
        self.print_reactions()
        self.generate_report() # [MODIFIED] 해석 완료 후 리포트 자동 생성

    def plot(self, factor=1.0, show_window=True, filename="fem_result.png"):
        """PyVista를 사용하여 변위 결과를 시각화하거나 파일로 저장합니다."""
        if self.u is None: return
        print("5. Plotting results...")
        
        cells = np.hstack([np.full((len(self.tetra10_conn), 1), 10), self.tetra10_conn])
        grid = pv.UnstructuredGrid(cells, [pv.CellType.QUADRATIC_TETRA] * len(self.tetra10_conn), self.points)
        disp_vectors = self.u.reshape(-1, self.pd)
        if not np.all(np.isfinite(disp_vectors)):
            print("Warning: Invalid displacement values found. Replacing with 0 for visualization.")
            disp_vectors = np.nan_to_num(disp_vectors)
        grid['Displacement Vectors'] = disp_vectors
        scalars = np.linalg.norm(disp_vectors, axis=1)
        grid['Magnitude'] = scalars
        print(f"   - Max displacement: {np.max(scalars):.6e} m")
        warped = grid.warp_by_vector('Displacement Vectors', factor=factor)
        
        plotter = pv.Plotter(off_screen=not show_window)
        plotter.add_axes()
        plotter.add_text(f"Deformed Shape (Magnitude) | Warp Factor: {factor}x", font_size=15)
        plotter.add_mesh(warped, scalars='Magnitude', cmap='jet', show_edges=True)
        fixed_node_indices = np.unique(self.fixed_dofs // self.pd)
        plotter.add_points(self.points[fixed_node_indices], color='blue', point_size=10, render_points_as_spheres=True, label='Fixed Nodes')

        if self.reaction_forces is not None:
            bounds = grid.bounds
            diag_length = np.sqrt(np.sum((np.array(bounds[1::2]) - np.array(bounds[::2]))**2))
            label_offset = diag_length * 0.05 if diag_length > 1e-6 else 0.05
            for info in self.fixed_nodes_info:
                pos = info['pos'].copy(); pos[1] += label_offset
                reactions = self.reaction_forces[3 * info['node_idx'] : 3 * info['node_idx'] + 3]
                label = f"Rx:{np.nan_to_num(reactions[0]):.2e}\nRy:{np.nan_to_num(reactions[1]):.2e}\nRz:{np.nan_to_num(reactions[2]):.2e}"
                plotter.add_point_labels(pos, [label], font_size=10, always_visible=True)
        
        plotter.add_legend()
        if show_window:
            plotter.show()
        else:
            plotter.camera_position = 'iso'
            plotter.screenshot(filename, window_size=[1024, 768])
            print(f"   - Result view saved to {filename}")

    def generate_report(self, filename="FEM_Report.docx"):
        """[NEW] Microsoft Word(.docx) 형식으로 해석 리포트를 생성합니다."""
        if not DOCX_AVAILABLE:
            print("\nWarning: python-docx is not installed. Skipping report generation.")
            print("Please run 'pip install python-docx' to enable this feature.")
            return
            
        print(f"\n6. Generating analysis report to {filename}...")
        report_generator = ReportGenerator(self)
        report_generator.create_report(filename)
        print("   - Report generation complete.")


class ReportGenerator:
    """[NEW] FEM 해석 결과를 Microsoft Word 문서로 생성하는 클래스."""
    def __init__(self, analysis_instance: ForceAnalysis):
        self.analysis = analysis_instance
        self.doc = Document()

    def create_report(self, filename):
        """지정된 파일 이름으로 전체 리포트를 생성합니다."""
        self.doc.add_heading('Finite Element Analysis Report', 0)

        self._add_results_image()
        self._add_analysis_parameters() # [MODIFIED] 해석 파라미터 섹션 추가
        self._add_mesh_info()
        self._add_boundary_conditions()
        self._add_mesh_quality()
        self._add_reaction_forces()

        try:
            self.doc.save(filename)
            print(f"   - Report successfully saved to {os.path.abspath(filename)}")
        except Exception as e:
            print(f"Error saving report file: {e}")

    def _add_results_image(self):
        """해석 결과의 등각 투영 뷰 이미지를 문서에 추가합니다."""
        self.doc.add_heading('Results Overview (Isometric View)', level=1)
        image_path = 'iso_view_report.png'
        self.analysis.plot(factor=1.0, show_window=False, filename=image_path)

        if os.path.exists(image_path):
            self.doc.add_paragraph('Deformed shape with displacement magnitude.')
            self.doc.add_picture(image_path, width=Inches(6.0))
        else:
            self.doc.add_paragraph('Could not generate result image.')
            
    def _add_analysis_parameters(self):
        """[NEW] 해석에 사용된 파라미터(재료 물성치)를 문서에 추가합니다."""
        self.doc.add_heading('Analysis Parameters', level=1)
        p = self.doc.add_paragraph()
        p.add_run('Young\'s Modulus (E): ').bold = True
        p.add_run(f"{self.analysis.E:.2e} Pa\n")
        p.add_run('Poisson\'s Ratio (v): ').bold = True
        p.add_run(f"{self.analysis.v}")

    def _add_mesh_info(self):
        """메쉬의 절점 및 요소 정보를 문서에 추가합니다."""
        self.doc.add_heading('Mesh Information', level=1)
        p = self.doc.add_paragraph()
        p.add_run('Total Nodes: ').bold = True
        p.add_run(f"{self.analysis.num_nodes}\n")
        p.add_run('Element Type: ').bold = True
        p.add_run("10-Node Quadratic Tetrahedron (Tetra10)\n")
        p.add_run('Total Elements: ').bold = True
        p.add_run(f"{len(self.analysis.tetra10_conn)}")

    def _add_boundary_conditions(self):
        """적용된 하중 및 고정 조건을 문서에 추가합니다."""
        self.doc.add_heading('Boundary Conditions', level=1)
        
        self.doc.add_heading('Applied Loads', level=2)
        force_info = self.analysis.force_data
        force_vec = (force_info['force_x'], force_info['force_y'], force_info['force_z'])
        force_pos = (force_info['force_x_pstn'], force_info['force_y_pstn'], force_info['force_z_pstn'])
        self.doc.add_paragraph(f" - Force Vector (Fx, Fy, Fz): {force_vec} N")
        self.doc.add_paragraph(f" - Application Point (x, y, z): {force_pos} m")

        self.doc.add_heading('Fixed Supports (Constraints)', level=2)
        for i, fix in enumerate(self.analysis.fix_data):
            pos = (fix['pos_x'], fix['pos_y'], fix['pos_z'])
            fixed_dofs = [d for d, c in zip(['X', 'Y', 'Z'], [fix['fix_x'], fix['fix_y'], fix['fix_z']]) if c == 0]
            self.doc.add_paragraph(f" - Fix Point {i+1} at {pos}: Constrained DOFs [{', '.join(fixed_dofs)}]")

    def _add_mesh_quality(self):
        """메쉬 품질(자코비안 행렬식)에 대한 정보를 추가합니다."""
        self.doc.add_heading('Mesh Quality Check', level=1)
        count = self.analysis.negative_detJ_count
        if count > 0:
            self.doc.add_paragraph(f" - Warning: Found {count} integration points with a non-positive Jacobian determinant (detJ <= 0). This indicates distorted elements which can lead to inaccurate results.")
        else:
            self.doc.add_paragraph(" - All elements passed the Jacobian determinant check (all detJ > 0).")

    def _add_reaction_forces(self):
        """계산된 반력을 테이블 형태로 문서에 추가합니다."""
        self.doc.add_heading('Reaction Force Results', level=1)

        if self.analysis.reaction_forces is None:
            self.doc.add_paragraph("Reaction forces were not calculated.")
            return
        
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Fix Point'
        hdr_cells[1].text = 'Node ID'
        hdr_cells[2].text = 'Rx (N)'
        hdr_cells[3].text = 'Ry (N)'
        hdr_cells[4].text = 'Rz (N)'

        total_reaction = np.zeros(3)
        for i, info in enumerate(self.analysis.fixed_nodes_info):
            node_idx = info['node_idx']
            reactions = self.analysis.reaction_forces[3 * node_idx : 3 * node_idx + 3]
            total_reaction += reactions
            row_cells = table.add_row().cells
            row_cells[0].text = str(i + 1)
            row_cells[1].text = str(node_idx)
            row_cells[2].text = f"{reactions[0]:.4e}"
            row_cells[3].text = f"{reactions[1]:.4e}"
            row_cells[4].text = f"{reactions[2]:.4e}"

        row_cells = table.add_row().cells
        row_cells[0].text = 'Total Reaction'
        row_cells[0].merge(row_cells[1])
        row_cells[2].text = f"{total_reaction[0]:.4e}"
        row_cells[3].text = f"{total_reaction[1]:.4e}"
        row_cells[4].text = f"{total_reaction[2]:.4e}"
        
        self.doc.add_paragraph("\nNote: For static equilibrium, the 'Total Reaction' should be equal and opposite to the sum of applied forces.")

