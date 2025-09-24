import sys
import numpy as np
import matplotlib.pyplot as plt
import meshio
from mpl_toolkits.mplot3d import Axes3D
import tkinter as tk
from tkinter import filedialog
from datetime import datetime
import scipy.linalg as la
import io

from PyQt5 import uic
from PyQt5.QtWidgets import QApplication, QDialog, QWidget, QVBoxLayout, QLabel, QMessageBox, QInputDialog
from PyQt5.QtCore import QStringListModel
from PyQt5.QtGui import QPixmap

# sectionproperties 라이브러리 임포트
import sectionproperties.pre.library.steel_sections as lib_s
import sectionproperties.pre.library.primitive_sections as lib_p
import sectionproperties.pre.pre as pre_
from sectionproperties.analysis.section import Section

# --- UI 파일 로딩 ---
try:
    Ui_Dialog, QtBaseClass = uic.loadUiType('Beam_analysis.ui')
except FileNotFoundError as e:
    print(f"Error: Could not find a required UI file: {e.filename}")
    print("Please make sure all .ui files are present in the same directory.")
    sys.exit(1)

# --- 단면 특성 계산 함수 (스케일링 및 메시 문제 해결) ---
def calculate_section_properties(section_type, params, rotate=False):
    """
    sectionproperties 라이브러리를 사용하여 단면 특성을 계산합니다.
    작은 값으로 인한 수치 불안정성을 피하기 위해 내부적으로 단위를 스케일링합니다.
    """
    geometry = None
    scale_factor = 1000.0
    scaled_params = {k: v * scale_factor for k, v in params.items()}
    try:
        if section_type == "hollow box section":
            geometry = lib_s.rectangular_hollow_section(d=scaled_params['d'], b=scaled_params['b'], t=scaled_params['t'], r_out=scaled_params.get('r_out', 0), n_r=8)
        elif section_type == "I section":
            geometry = lib_s.i_section(d=scaled_params['d'], b=scaled_params['b'], t_f=scaled_params['t_f'], t_w=scaled_params['t_w'], r=scaled_params.get('r', 0), n_r=8)
        elif section_type == "C section":
            geometry = lib_s.channel_section(d=scaled_params['d'], b=scaled_params['b'], t_f=scaled_params['t_f'], t_w=scaled_params['t_w'], r=scaled_params.get('r', 0), n_r=8)
        elif section_type == "L section":
            geometry = lib_s.angle_section(d=scaled_params['d'], b=scaled_params['b'], t=scaled_params['t'], r_r=scaled_params.get('r_r', 0), r_t=scaled_params.get('r_t', 0), n_r=8)
        elif section_type == "rectangular section":
            geometry = lib_p.rectangular_section(d=scaled_params['d'], b=scaled_params['b'])
        elif section_type == "circular section":
            geometry = lib_p.circular_section(d=scaled_params['d'], n=64)
        elif section_type == "hollow circular section":
            geometry = lib_s.circular_hollow_section(d=scaled_params['d'], t=scaled_params['t'], n=64)
        else:
            print(f"Warning: Unknown section type '{section_type}'.")
            return (0,) * 8
        t_vals = [v for k, v in scaled_params.items() if 't' in k and v > 0]
        if not t_vals:
            dim_vals = [v for k, v in scaled_params.items() if k in ['d', 'b'] and v > 0]
            mesh_size_ref = min(dim_vals) if dim_vals else 1.0
        else:
            mesh_size_ref = min(t_vals)
        mesh_size = mesh_size_ref / 10.0
        geometry.create_mesh(mesh_sizes=[mesh_size])
        section = Section(geometry, time_info=False)
        section.calculate_geometric_properties()
        section.calculate_warping_properties()
        A_mm2 = section.get_area(); J_mm4 = section.get_j(); (ixx_c_mm4, iyy_c_mm4, ixy_c_mm4) = section.get_ic()
        (Asx_mm2, Asy_mm2) = section.get_as(); (cx_mm, cy_mm) = section.get_c()
        vertices_mm = section.mesh['vertices']
        c_y_max_mm = max(abs(vertices_mm[:, 0] - cx_mm)); c_z_max_mm = max(abs(vertices_mm[:, 1] - cy_mm))
        A = A_mm2 / (scale_factor ** 2); J = J_mm4 / (scale_factor ** 4); I_y = ixx_c_mm4 / (scale_factor ** 4); I_z = iyy_c_mm4 / (scale_factor ** 4)
        kappa_y = Asx_mm2 / A_mm2 if A_mm2 > 0 else 0; kappa_z = Asy_mm2 / A_mm2 if A_mm2 > 0 else 0
        c_y_max = c_y_max_mm / scale_factor; c_z_max = c_z_max_mm / scale_factor
        if rotate:
            I_y, I_z = I_z, I_y; kappa_y, kappa_z = kappa_z, kappa_y; c_y_max, c_z_max = c_z_max, c_y_max
        I_x_final, I_y_final = I_y, I_z
        return A, I_x_final, I_y_final, J, kappa_y, kappa_z, c_y_max, c_z_max
    except Exception as e:
        print(f"Error in sectionproperties for {section_type} with SCALED params {scaled_params}: {e}")
        return (0,) * 8

# --- 단면 정보 입력을 위한 다이얼로그 클래스들 ---
class BaseSectionDialog(QDialog):
    """단면 다이얼로그의 기본 클래스"""
    def __init__(self, ui_file, parent=None, initial_data=None):
        super().__init__(parent)
        try:
            uic.loadUi(ui_file, self)
            self.setWindowTitle(f"Section Information")
            if initial_data: self.set_data(initial_data)
        except FileNotFoundError:
            self.setWindowTitle("UI File Error"); self.setLayout(QVBoxLayout()); self.layout().addWidget(QLabel(f"Could not load {ui_file}"))
    def get_values(self): raise NotImplementedError
    def set_data(self, data): pass
    def _get_float(self, w, default=0.0): return float(getattr(self, w).text().strip()) if hasattr(self, w) and getattr(self, w).text().strip() else default
    def _set_text(self, w, v):
        if hasattr(self, w) and v is not None: getattr(self, w).setText(str(v))
    def _get_bool(self, w, default=False): return getattr(self, w).isChecked() if hasattr(self, w) else default
    def _set_check(self, w, v):
        if hasattr(self, w) and v is not None: getattr(self, w).setChecked(bool(v))

class ISectionDialog(BaseSectionDialog):
    def __init__(self, parent=None, initial_data=None): super().__init__('I_section.ui', parent, initial_data)
    def get_values(self): return {'d': self._get_float('d_input'), 'b': self._get_float('b_input'), 't_w': self._get_float('tw_input'), 't_f': self._get_float('tf_input'), 'r': self._get_float('r_input'), 'rotate': self._get_bool('rotate_check')}
    def set_data(self, data): self._set_text('d_input', data.get('d')); self._set_text('b_input', data.get('b')); self._set_text('tw_input', data.get('t_w')); self._set_text('tf_input', data.get('t_f')); self._set_text('r_input', data.get('r')); self._set_check('rotate_check', data.get('rotate'))

class CSectionDialog(BaseSectionDialog):
    def __init__(self, parent=None, initial_data=None): super().__init__('C_section.ui', parent, initial_data)
    def get_values(self): return {'d': self._get_float('d_input'), 'b': self._get_float('b_input'), 't_f': self._get_float('tf_input'), 't_w': self._get_float('tw_input'), 'r': self._get_float('r_input'), 'rotate': self._get_bool('rotate_check')}
    def set_data(self, data): self._set_text('d_input', data.get('d')); self._set_text('b_input', data.get('b')); self._set_text('tf_input', data.get('t_f')); self._set_text('tw_input', data.get('t_w')); self._set_text('r_input', data.get('r')); self._set_check('rotate_check', data.get('rotate'))

class LSectionDialog(BaseSectionDialog):
    def __init__(self, parent=None, initial_data=None): super().__init__('L_section.ui', parent, initial_data)
    def get_values(self): return {'d': self._get_float('d_input'), 'b': self._get_float('b_input'), 't': self._get_float('t_input'), 'r_r': self._get_float('r_r_input'), 'r_t': self._get_float('r_t_input'), 'rotate': self._get_bool('rotate_check')}
    def set_data(self, data): self._set_text('d_input', data.get('d')); self._set_text('b_input', data.get('b')); self._set_text('t_input', data.get('t')); self._set_text('r_r_input', data.get('r_r')); self._set_text('r_t_input', data.get('r_t')); self._set_check('rotate_check', data.get('rotate'))

class HollowBoxSectionDialog(BaseSectionDialog):
    def __init__(self, parent=None, initial_data=None): super().__init__('hollow_box_section.ui', parent, initial_data)
    def get_values(self): return {'d': self._get_float('d_input'), 'b': self._get_float('b_input'), 't': self._get_float('t_input'), 'r_out': self._get_float('r_input'), 'rotate': self._get_bool('rotate_check')}
    def set_data(self, data): self._set_text('d_input', data.get('d')); self._set_text('b_input', data.get('b')); self._set_text('t_input', data.get('t')); self._set_text('r_input', data.get('r_out')); self._set_check('rotate_check', data.get('rotate'))

class RectangularSectionDialog(BaseSectionDialog):
    def __init__(self, parent=None, initial_data=None): super().__init__('Rectangular_section.ui', parent, initial_data)
    def get_values(self): return {'d': self._get_float('d_input'), 'b': self._get_float('b_input'), 'rotate': self._get_bool('rotate_check')}
    def set_data(self, data): self._set_text('d_input', data.get('d')); self._set_text('b_input', data.get('b')); self._set_check('rotate_check', data.get('rotate'))

class CircularSectionDialog(BaseSectionDialog):
    def __init__(self, parent=None, initial_data=None): super().__init__('circular_section.ui', parent, initial_data)
    def get_values(self): return {'d': self._get_float('d_input')}
    def set_data(self, data): self._set_text('d_input', data.get('d'))

class HollowCircularSectionDialog(BaseSectionDialog):
    def __init__(self, parent=None, initial_data=None): super().__init__('hollow_circular_section.ui', parent, initial_data)
    def get_values(self): return {'d': self._get_float('d_input'), 't': self._get_float('t_input')}
    def set_data(self, data): self._set_text('d_input', data.get('d')); self._set_text('t_input', data.get('t'))

class BoundaryConditionDialog(QDialog):
    def __init__(self, parent=None, initial_data=None):
        super().__init__(parent)
        try:
            uic.loadUi('beam_fix.ui', self); self.setWindowTitle("Set Boundary Condition")
            if hasattr(self, 'fix_radioButton'): self.fix_radioButton.toggled.connect(self._update_widget_states)
            if hasattr(self, 'applyForce_radioButton'): self.applyForce_radioButton.toggled.connect(self._update_widget_states)
            if initial_data: self.set_data(initial_data)
            else: self._update_widget_states()
        except FileNotFoundError:
            self.setWindowTitle("Error"); self.setLayout(QVBoxLayout()); self.layout().addWidget(QLabel("Could not load beam_fix.ui"))
    def _update_widget_states(self):
        is_fix = self.fix_radioButton.isChecked(); is_force = self.applyForce_radioButton.isChecked()
        for name in ['x_checkBox','y_checkBox','z_checkBox','rx_checkBox','ry_checkBox','rz_checkBox']:
            if hasattr(self, name): getattr(self, name).setEnabled(is_fix)
        for name in ['force_x_input', 'force_y_input', 'force_z_input']:
            if hasattr(self, name): getattr(self, name).setEnabled(is_force)
    def set_data(self, data):
        if data.get('type') == 'Fix':
            self.fix_radioButton.setChecked(True); self.x_checkBox.setChecked(data.get('fix_x', False)); self.y_checkBox.setChecked(data.get('fix_y', False)); self.z_checkBox.setChecked(data.get('fix_z', False)); 
            if hasattr(self, 'rx_checkBox'): self.rx_checkBox.setChecked(data.get('fix_rx', False))
            if hasattr(self, 'ry_checkBox'): self.ry_checkBox.setChecked(data.get('fix_ry', False))
            if hasattr(self, 'rz_checkBox'): self.rz_checkBox.setChecked(data.get('fix_rz', False))
        elif data.get('type') == 'Force':
            self.applyForce_radioButton.setChecked(True); self.force_x_input.setText(str(data.get('force_x', 0.0))); self.force_y_input.setText(str(data.get('force_y', 0.0))); self.force_z_input.setText(str(data.get('force_z', 0.0)))
        self._update_widget_states()
    def get_data(self):
        try:
            if self.fix_radioButton.isChecked():
                return {'type': 'Fix', 'fix_x': self.x_checkBox.isChecked(), 'fix_y': self.y_checkBox.isChecked(), 'fix_z': self.z_checkBox.isChecked(), 'fix_rx': self.rx_checkBox.isChecked() if hasattr(self, 'rx_checkBox') else False, 'fix_ry': self.ry_checkBox.isChecked() if hasattr(self, 'ry_checkBox') else False, 'fix_rz': self.rz_checkBox.isChecked() if hasattr(self, 'rz_checkBox') else False}
            elif self.applyForce_radioButton.isChecked():
                return {'type': 'Force', 'force_x': float(self.force_x_input.text()), 'force_y': float(self.force_y_input.text()), 'force_z': float(self.force_z_input.text())}
            return None
        except ValueError:
            QMessageBox.critical(self, "Input Error", "Please enter valid numbers for forces."); return None

# --- 메인 윈도우 클래스 ---
class BeamAnalysisWindow(QDialog, Ui_Dialog):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.setWindowTitle("3D Timoshenko Beam Analysis")

        self.mesh = None; self.points = None; self.conn = None
        self.u = None; self.smoothed_stresses = None
        self.natural_frequencies = None; self.mode_shapes = None
        self.section_data = []; self.bc_data = []
        self.static_plot_buffer = None; self.modal_plot_buffers = []

        self.list_model = QStringListModel()
        self.listView.setModel(self.list_model)

        self.section_type_combo.addItems(["I section", "C section", "L section", "hollow box section",
                                          "rectangular section", "circular section", "hollow circular section"])

        self.mesh_sel_btn.clicked.connect(self.mesh_load)
        self.beam_button.clicked.connect(self.assign_beam_section)
        self.bc_button.clicked.connect(self.assign_bc)
        self.edit_button.clicked.connect(self.edit_item)
        self.remove_button.clicked.connect(self.remove_item)
        self.mesh_update_button.clicked.connect(self.mesh_update)
        self.run_button.clicked.connect(self.run_simulation)
        self.plot_button.clicked.connect(self.plot_static_results)
        
        # UI 파일에 plot_modal_button이 있다고 가정
        if hasattr(self, 'plot_modal_button'):
            self.plot_modal_button.clicked.connect(self.plot_modal_results)

    def mesh_load(self):
        root = tk.Tk(); root.withdraw()
        mesh_path = filedialog.askopenfilename(title="Select Gmsh .msh file", filetypes=[("Gmsh mesh", "*.msh")])
        if not mesh_path: return
        try:
            self.mesh = meshio.read(mesh_path); self.root_indicate.setText(mesh_path)
            self.points = self.mesh.points; self.conn = self.mesh.cells_dict.get('line')
            if self.conn is None: raise ValueError("No 'line' elements in .msh file.")
            physical_groups = list(self.mesh.field_data.keys()) if self.mesh.field_data else []
            self.physical_group_combo.clear(); self.bc_combo.clear()
            self.physical_group_combo.addItems(physical_groups); self.bc_combo.addItems(physical_groups)
            self.section_data.clear(); self.bc_data.clear(); self.update_list_view()
        except Exception as e:
            QMessageBox.critical(self, "Mesh Load Error", f"Failed to read mesh file:\n{e}"); self.mesh = None

    def assign_beam_section(self):
        group = self.physical_group_combo.currentText()
        if not group: QMessageBox.warning(self, "Selection Error", "Please select a physical group."); return
        if any(d['group'] == group for d in self.section_data):
            if QMessageBox.question(self, 'Confirm Overwrite', f"Overwrite existing section for '{group}'?", QMessageBox.Yes | QMessageBox.No) == QMessageBox.No: return
            self.section_data = [d for d in self.section_data if d['group'] != group]
        section_type = self.section_type_combo.currentText()
        dialog_map = {"I section": ISectionDialog, "C section": CSectionDialog, "L section": LSectionDialog, "hollow box section": HollowBoxSectionDialog, "rectangular section": RectangularSectionDialog, "circular section": CircularSectionDialog, "hollow circular section": HollowCircularSectionDialog}
        dialog_class = dialog_map.get(section_type)
        if not dialog_class: QMessageBox.critical(self, "Error", f"No dialog found for section type '{section_type}'."); return
        dialog = dialog_class(self)
        if dialog.exec_() == QDialog.Accepted:
            params = dialog.get_values()
            if params is not None:
                is_rotated = params.pop('rotate', False)
                self.section_data.append({'group': group, 'type': section_type, 'params': params, 'rotate': is_rotated})
                self.update_list_view()

    def assign_bc(self):
        group = self.bc_combo.currentText()
        if not group: QMessageBox.warning(self, "Selection Error", "Please select a physical group for the BC."); return
        if any(d['group'] == group for d in self.bc_data):
            if QMessageBox.question(self, 'Confirm Overwrite', f"Overwrite existing BC for '{group}'?", QMessageBox.Yes | QMessageBox.No) == QMessageBox.No: return
            self.bc_data = [d for d in self.bc_data if d['group'] != group]
        dialog = BoundaryConditionDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            data = dialog.get_data()
            if data:
                data['group'] = group; self.bc_data.append(data); self.update_list_view()

    def edit_item(self):
        selected = self.listView.selectedIndexes()
        if not selected: QMessageBox.warning(self, "Selection Error", "Please select an item to edit."); return
        row = selected[0].row(); num_sections = len(self.section_data)
        if row < num_sections:
            data = self.section_data[row]; section_type = data['type']; initial_params = {**data['params'], 'rotate': data.get('rotate', False)}
            dialog_map = {"I section": ISectionDialog, "C section": CSectionDialog, "L section": LSectionDialog, "hollow box section": HollowBoxSectionDialog, "rectangular section": RectangularSectionDialog, "circular section": CircularSectionDialog, "hollow circular section": HollowCircularSectionDialog}
            dialog_class = dialog_map.get(section_type)
            if not dialog_class: return
            dialog = dialog_class(self, initial_data=initial_params)
            if dialog.exec_() == QDialog.Accepted:
                new_params = dialog.get_values()
                if new_params: data['rotate'] = new_params.pop('rotate', False); data['params'] = new_params
        else:
            bc_index = row - num_sections; data = self.bc_data[bc_index]
            dialog = BoundaryConditionDialog(self, initial_data=data)
            if dialog.exec_() == QDialog.Accepted:
                new_data = dialog.get_data()
                if new_data: new_data['group'] = data['group']; self.bc_data[bc_index] = new_data
        self.update_list_view()

    def remove_item(self):
        selected = self.listView.selectedIndexes()
        if not selected: QMessageBox.warning(self, "Selection Error", "Please select an item to remove."); return
        row = selected[0].row(); num_sections = len(self.section_data)
        item_text = self.section_data[row]['group'] if row < num_sections else self.bc_data[row - num_sections]['group']
        if QMessageBox.question(self, 'Confirm Remove', f"Remove item for '{item_text}'?", QMessageBox.Yes | QMessageBox.No) == QMessageBox.Yes:
            if row < num_sections: self.section_data.pop(row)
            else: self.bc_data.pop(row - num_sections)
            self.update_list_view()

    def update_list_view(self):
        display_list = []
        for item in self.section_data:
            params_str = ", ".join([f"{k}={v}" for k, v in item['params'].items()])
            rotate_str = ", rotated" if item.get('rotate') else ""
            display_list.append(f"[Section] {item['group']}: {item['type']}{rotate_str}, {params_str}")
        for item in self.bc_data:
            details = ""
            if item['type'] == 'Fix':
                fixes = [f for f in ['X','Y','Z','RX','RY','RZ'] if item.get(f'fix_{f.lower()}')]
                details = f"Fix ({', '.join(fixes) or 'None'})"
            elif item['type'] == 'Force':
                forces = f"F=({item.get('force_x', 0)}, {item.get('force_y', 0)}, {item.get('force_z', 0)})"
                details = f"Force {forces}"
            display_list.append(f"[BC] {item['group']}: {details}")
        self.list_model.setStringList(display_list)

    def mesh_update(self): 
        if not self.mesh or not self.section_data: QMessageBox.warning(self, "No Data", "Please load a mesh and assign sections first."); return
        import pyvista as pv
        plotter = pv.Plotter()
        try: E = float(self.young_input.text()); nu = float(self.poisson_input.text())
        except: QMessageBox.warning(self, "No Data", "Please input material information first."); return
        try: Material_ = pre_.Material(name='default', elastic_modulus=E, poissons_ratio=nu, yield_strength=1, density=1, color='w')
        except: QMessageBox.warning(self, "No Data", "Material assign error."); return
        props_map = {}
        for sec_data in self.section_data:
            group = sec_data['group']; section_type = sec_data['type']; params = sec_data['params']; geometry = None
            if section_type == "hollow box section": geometry = lib_s.rectangular_hollow_section(d=params['d'], b=params['b'], t=params['t'], r_out=params.get('r_out', 0), n_r=8)
            elif section_type == "I section": geometry = lib_s.i_section(d=params['d'], b=params['b'], t_f=params['t_f'], t_w=params['t_w'], r=params.get('r', 0), n_r=8)
            elif section_type == "C section": geometry = lib_s.channel_section(d=params['d'], b=params['b'], t_f=params['t_f'], t_w=params['t_w'], r=params.get('r', 0), n_r=8)
            elif section_type == "L section": geometry = lib_s.angle_section(d=params['d'], b=params['b'], t=params['t'], r_r=params.get('r_r', 0), r_t=params.get('r_t', 0), n_r=8)
            elif section_type == "rectangular section": geometry = lib_p.rectangular_section(d=params['d'], b=params['b'], material=Material_)
            elif section_type == "circular section": geometry = lib_p.circular_section(d=params['d'], material=Material_, n=64)
            elif section_type == "hollow circular section": geometry = lib_s.circular_hollow_section(d=params['d'], t=params['t'], n=64)
            if geometry is None: continue
            t_vals = [v for k, v in params.items() if 't' in k and v > 0]; mesh_size_ref = min(t_vals) if t_vals else params.get('d', 1)
            mesh_size = mesh_size_ref / 4.0 if mesh_size_ref > 0 else 0.1; geometry.create_mesh(mesh_sizes=[mesh_size])
            if sec_data.get('rotate', False):
                vertices = geometry.mesh['vertices']; rotated_vertices = np.zeros_like(vertices)
                rotated_vertices[:, 0] = -vertices[:, 1]; rotated_vertices[:, 1] = vertices[:, 0]
                geometry.mesh['vertices'] = rotated_vertices
            props_map[group] = geometry
        group_id_to_name = {v[0]: k for k, v in self.mesh.field_data.items()}; line_phys_tags = self.mesh.cell_data_dict['gmsh:physical']['line']; eps = 1e-6
        for i, element in enumerate(self.conn):
            phys_tag = line_phys_tags[i]; group_name = group_id_to_name.get(phys_tag)
            if group_name not in props_map: continue
            geometry = props_map[group_name]; vertices_2d = geometry.mesh['vertices']; triangles = geometry.mesh['triangles']
            p1 = self.points[element[0]]; p2 = self.points[element[1]]; L = np.linalg.norm(p2 - p1)
            if L == 0: continue
            dir_vec = (p2 - p1) / L; Cxx, Cyx, Czx = dir_vec
            if Cxx**2 + Cyx**2 < eps**2: sign = 1 if Czx > 0 else -1; lambda_mat = np.array([[0., 0., sign], [0., 1., 0.], [-sign, 0., 0.]])
            else: D = np.sqrt(Cxx**2 + Cyx**2); lambda_mat = np.array([[Cxx, Cyx, Czx], [-Cyx/D, Cxx/D, 0.], [-Cxx*Czx/D, -Cyx*Czx/D, D]])
            local_points = np.zeros((vertices_2d.shape[0], 3)); local_points[:,1] = vertices_2d[:,0]; local_points[:,2] = vertices_2d[:,1]
            global_dirs = lambda_mat @ local_points.T
            global_points_start = (p1[:, None] + global_dirs).T; global_points_end = (p2[:, None] + global_dirs).T
            all_points = np.vstack((global_points_start, global_points_end)); num_pts = len(global_points_start)
            cells = [[6, tri[0], tri[1], tri[2], tri[0]+num_pts, tri[1]+num_pts, tri[2]+num_pts] for tri in triangles]
            grid = pv.UnstructuredGrid(np.hstack(cells), np.full(len(triangles), pv._vtk.VTK_WEDGE, dtype=np.uint8), all_points)
            plotter.add_mesh(grid, opacity=1, show_edges=True)
        plotter.add_axes(); plotter.show()

    def run_simulation(self):
        if not self.mesh:
            QMessageBox.warning(self, "Error", "Please load a mesh file first.")
            return
        try:
            E = float(self.young_input.text())
            nu = float(self.poisson_input.text())
            G = E / (2 * (1 + nu))
            pd = 6
            num_nodes = len(self.points)

            props_map = {sec['group']: calculate_section_properties(sec['type'], sec['params'], sec.get('rotate', False)) for sec in self.section_data}
            group_id_to_name = {v[0]: k for k, v in self.mesh.field_data.items()}
            line_phys_tags = self.mesh.cell_data_dict['gmsh:physical']['line']

            k_global = np.zeros((pd * num_nodes, pd * num_nodes))
            m_global = np.zeros((pd * num_nodes, pd * num_nodes))
            eps = 1e-6

            for i, element in enumerate(self.conn):
                phys_tag = line_phys_tags[i]
                group_name = group_id_to_name.get(phys_tag)
                if not group_name or group_name not in props_map:
                    QMessageBox.critical(self, "Error", f"Section properties not defined for physical group '{group_name}'.")
                    return
                
                A, I_x, I_y, J, kappa_y, kappa_z, c_y_max, c_z_max = props_map[group_name]
                p1, p2 = self.points[element[0]], self.points[element[1]]
                L = np.linalg.norm(p2 - p1)
                
                k_ = self.get_timoshenko_stiffness_matrix(L, E, G, A, I_x, I_y, J, kappa_y, kappa_z)
                m_ = self.get_lumped_mass_matrix(L, A, I_x, I_y, J, 7850)

                direction = (p2 - p1) / L
                Cxx, Cyx, Czx = direction
                if Cxx**2 + Cyx**2 < eps**2:
                    lambda_ = np.array([[0., 0., 1. if Czx > 0 else -1.], [0., 1., 0.], [-1. if Czx > 0 else 1., 0., 0.]])
                else:
                    D = np.sqrt(Cxx**2 + Cyx**2)
                    lambda_ = np.array([[Cxx, Cyx, Czx], [-Cyx/D, Cxx/D, 0], [-Cxx*Czx/D, -Cyx*Czx/D, D]])
                
                R_theta = np.kron(np.eye(4, dtype=float), lambda_)
                k_local = R_theta.T @ k_ @ R_theta
                m_local = R_theta.T @ m_ @ R_theta

                for j, J_node in enumerate(element):
                    for l, L_node in enumerate(element):
                        k_global[6*J_node:6*J_node+6, 6*L_node:6*L_node+6] += k_local[6*j:6*j+6, 6*l:6*l+6]
                        m_global[6*J_node:6*J_node+6, 6*L_node:6*L_node+6] += m_local[6*j:6*j+6, 6*l:6*l+6]

            self.u = np.zeros(pd * num_nodes); f = np.zeros(pd * num_nodes); Up_node = []
            for bc in self.bc_data:
                nodes_to_apply = self.bc_nodes_indexing('vertex', bc['group'])
                for num in nodes_to_apply:
                    if bc['type'] == 'Fix':
                        if bc.get('fix_x'): Up_node.append(6*num + 0)
                        if bc.get('fix_y'): Up_node.append(6*num + 1)
                        if bc.get('fix_z'): Up_node.append(6*num + 2)
                        if bc.get('fix_rx'): Up_node.append(6*num + 3)
                        if bc.get('fix_ry'): Up_node.append(6*num + 4)
                        if bc.get('fix_rz'): Up_node.append(6*num + 5)
                    elif bc['type'] == 'Force':
                        f[6*num + 0] += bc.get('force_x', 0); f[6*num + 1] += bc.get('force_y', 0); f[6*num + 2] += bc.get('force_z', 0)
            
            Up_node = sorted(list(set(Up_node)))
            Fp_node = [x for x in range(pd*num_nodes) if x not in Up_node]

            u_temp = np.zeros_like(self.u)
            for dof in Up_node: u_temp[dof] = 0.0
            k_ff = k_global[np.ix_(Fp_node, Fp_node)]
            k_fs = k_global[np.ix_(Fp_node, Up_node)]
            f_f = f[Fp_node] - k_fs @ u_temp[Up_node]
            u_temp[Fp_node] = np.linalg.solve(k_ff, f_f)
            self.u = u_temp

            nodal_stresses = np.zeros(num_nodes); node_contributions = np.zeros(num_nodes, dtype=int)
            for i, element in enumerate(self.conn):
                phys_tag = line_phys_tags[i]; group_name = group_id_to_name.get(phys_tag)
                A, I_x, I_y, J, kappa_y, kappa_z, c_y_max, c_z_max = props_map[group_name]
                node1_idx, node2_idx = element; p1, p2 = self.points[node1_idx], self.points[node2_idx]; L = np.linalg.norm(p2 - p1)
                k_ = self.get_timoshenko_stiffness_matrix(L, E, G, A, I_x, I_y, J, kappa_y, kappa_z)
                direction = (p2 - p1) / L; Cxx, Cyx, Czx = direction
                if Cxx**2 + Cyx**2 < eps**2: lambda_ = np.array([[0.,0.,1. if Czx>0 else -1.],[0.,1.,0.],[-1. if Czx>0 else 1.,0.,0.]])
                else: D = np.sqrt(Cxx**2+Cyx**2); lambda_ = np.array([[Cxx,Cyx,Czx],[-Cyx/D,Cxx/D,0],[-Cxx*Czx/D,-Cyx*Czx/D,D]])
                R_theta = np.kron(np.eye(4, dtype=float), lambda_)
                u_global_element = np.concatenate((self.u[6*node1_idx:6*node1_idx+6], self.u[6*node2_idx:6*node2_idx+6]))
                f_local = k_ @ (R_theta @ u_global_element)
                sigma_axial = f_local[6] / A if A > 0 else 0
                sigma_bending1 = abs(f_local[4] * c_z_max / I_x if I_x > 0 else 0) + abs(f_local[5] * c_y_max / I_y if I_y > 0 else 0)
                sigma_bending2 = abs(f_local[10] * c_z_max / I_x if I_x > 0 else 0) + abs(f_local[11] * c_y_max / I_y if I_y > 0 else 0)
                nodal_stresses[node1_idx] += sigma_axial + sigma_bending1
                nodal_stresses[node2_idx] += sigma_axial + sigma_bending2
                node_contributions[node1_idx] += 1; node_contributions[node2_idx] += 1
            self.smoothed_stresses = np.divide(nodal_stresses, node_contributions, where=node_contributions!=0)

            m_ff = m_global[np.ix_(Fp_node, Fp_node)]
            try:
                A_matrix = np.linalg.inv(m_ff) @ k_ff
            except np.linalg.LinAlgError:
                 QMessageBox.critical(self, "Modal Analysis Error", "Mass matrix is singular. Cannot perform modal analysis."); return
            
            eigenvalues, eigenvectors = self.qr_algorithm(A_matrix)
            
            valid_indices = eigenvalues > 1e-6
            eigenvalues = eigenvalues[valid_indices]
            eigenvectors = eigenvectors[:, valid_indices]
            self.natural_frequencies = np.sqrt(eigenvalues)
            
            full_mode_shapes = np.zeros((pd * num_nodes, len(self.natural_frequencies)))
            full_mode_shapes[Fp_node, :] = eigenvectors
            self.mode_shapes = full_mode_shapes

            QMessageBox.information(self, "Success", "Static and Modal analyses completed.")
            
            self._generate_plots_for_report()
            
            if QMessageBox.question(self, 'Create Report', "Would you like to save a report?", QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes) == QMessageBox.Yes:
                self.create_report()

        except Exception as e:
            QMessageBox.critical(self, "Simulation Error", f"An error occurred during analysis:\n{e}")

    def qr_algorithm(self, A, max_iter=1000, tol=1e-9):
        A_k = A.copy()
        n = A.shape[0]
        V = np.eye(n)
        for _ in range(max_iter):
            Q, R = np.linalg.qr(A_k)
            A_k_new = R @ Q
            V = V @ Q
            if np.allclose(np.diag(A_k), np.diag(A_k_new), atol=tol):
                break
            A_k = A_k_new
        eigenvalues = np.diag(A_k_new)
        eigenvectors = V
        sort_indices = np.argsort(eigenvalues)
        return eigenvalues[sort_indices], eigenvectors[:, sort_indices]

    def create_report(self):
        if self.u is None: QMessageBox.warning(self, "No Data", "Please run the simulation first."); return
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            QMessageBox.critical(self, "Missing Library", "'python-docx' is required. Please install it."); return
        
        root = tk.Tk(); root.withdraw()
        file_path = filedialog.asksaveasfilename(title="Save Report", defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not file_path: return

        document = Document()
        title = document.add_heading('3D Beam Analysis Report', level=1); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_date = document.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"); p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        document.add_heading('1. Model Information', level=2)
        p = document.add_paragraph()
        p.add_run('Mesh Details\n').bold = True
        p.add_run(f"  • Number of Nodes: {len(self.points)}\n")
        p.add_run(f"  • Number of Elements: {len(self.conn)}\n\n")
        p.add_run('Material Properties\n').bold = True
        try:
            E = float(self.young_input.text()); nu = float(self.poisson_input.text())
            p.add_run(f"  • Young's Modulus (E): {E:,.2e} Pa\n"); p.add_run(f"  • Poisson's Ratio (ν): {nu}\n\n")
        except ValueError: p.add_run("  Material properties not specified.\n\n")
        
        p.add_run('Boundary Conditions & Loads').bold = True
        bc_table = document.add_table(rows=1, cols=3); bc_table.style = 'Table Grid'
        hdr_cells = bc_table.rows[0].cells; hdr_cells[0].text = 'Group'; hdr_cells[1].text = 'Type'; hdr_cells[2].text = 'Details'
        for item in self.bc_data:
            row_cells = bc_table.add_row().cells
            row_cells[0].text = item['group']
            row_cells[1].text = item['type']
            details = ""
            if item['type'] == 'Fix':
                fixes = [f for f in ['X','Y','Z','RX','RY','RZ'] if item.get(f'fix_{f.lower()}')]
                details = f"Fixed DOFs: {', '.join(fixes) or 'None'}"
            elif item['type'] == 'Force':
                details = f"Fx={item.get('force_x', 0)}, Fy={item.get('force_y', 0)}, Fz={item.get('force_z', 0)}"
            row_cells[2].text = details

        document.add_heading('2. Static Analysis Results', level=2)
        if self.static_plot_buffer:
            document.add_paragraph("Deformation and Stress Contour Plot:")
            document.add_picture(self.static_plot_buffer, width=Inches(6.0))
        document.add_paragraph("Nodal displacement and stress results:")
        table = document.add_table(rows=1, cols=8); table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Node ID'; hdr_cells[1].text = 'X (m)'; hdr_cells[2].text = 'Y (m)'; hdr_cells[3].text = 'Z (m)'; hdr_cells[4].text = 'Disp X (m)'; hdr_cells[5].text = 'Disp Y (m)'; hdr_cells[6].text = 'Disp Z (m)'; hdr_cells[7].text = 'Stress (MPa)'
        pd = 6; disp_x = self.u[0::pd]; disp_y = self.u[1::pd]; disp_z = self.u[2::pd]
        for i in range(len(self.points)):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i); row_cells[1].text = f"{self.points[i, 0]:.4f}"; row_cells[2].text = f"{self.points[i, 1]:.4f}"; row_cells[3].text = f"{self.points[i, 2]:.4f}"
            row_cells[4].text = f"{disp_x[i]:.4e}"; row_cells[5].text = f"{disp_y[i]:.4e}"; row_cells[6].text = f"{disp_z[i]:.4e}"; row_cells[7].text = f"{self.smoothed_stresses[i] / 1e6:.4f}"

        if self.natural_frequencies is not None:
            document.add_heading('3. Modal Analysis Results', level=2)
            table_modal = document.add_table(rows=1, cols=3); table_modal.style = 'Table Grid'
            hdr_cells_modal = table_modal.rows[0].cells
            hdr_cells_modal[0].text = 'Mode'; hdr_cells_modal[1].text = 'Frequency (rad/s)'; hdr_cells_modal[2].text = 'Frequency (Hz)'
            for i, freq_rads in enumerate(self.natural_frequencies[:10]):
                row_cells = table_modal.add_row().cells
                row_cells[0].text = str(i+1); row_cells[1].text = f"{freq_rads:.4f}"; row_cells[2].text = f"{freq_rads / (2*np.pi):.4f}"
            
            if self.modal_plot_buffers:
                document.add_paragraph("\nMode Shape Plots:")
                for i, buffer in enumerate(self.modal_plot_buffers):
                    document.add_picture(buffer, width=Inches(5.5))
                    last_paragraph = document.paragraphs[-1] 
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        try:
            document.save(file_path)
            QMessageBox.information(self, "Success", f"Report successfully saved to:\n{file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Save Error", f"Failed to save the report:\n{e}")

    def _generate_plots_for_report(self):
        """보고서에 삽입할 플롯 이미지를 버퍼에 저장합니다."""
        self.static_plot_buffer = None
        fig_static = self._create_static_plot_figure()
        if fig_static:
            buf = io.BytesIO()
            fig_static.savefig(buf, format='png', dpi=150)
            buf.seek(0)
            self.static_plot_buffer = buf
            plt.close(fig_static)

        self.modal_plot_buffers = []
        if self.mode_shapes is not None:
            num_modes_to_report = min(self.mode_shapes.shape[1], 5) 
            for i in range(1, num_modes_to_report + 1):
                fig_modal = self._create_modal_plot_figure(i)
                if fig_modal:
                    buf = io.BytesIO()
                    fig_modal.savefig(buf, format='png', dpi=150)
                    buf.seek(0)
                    self.modal_plot_buffers.append(buf)
                    plt.close(fig_modal)

    def _create_static_plot_figure(self):
        """정적 해석 결과를 Matplotlib Figure 객체로 생성하여 반환합니다."""
        if self.u is None: return None
        pd=6
        try: scale_factor = float(self.scale_factor_input.text())
        except (AttributeError, ValueError): scale_factor = 1.0
        x_orig, y_orig, z_orig = self.points[:,0], self.points[:,1], self.points[:,2]
        disp_x = self.u[0::pd]; disp_y = self.u[1::pd]; disp_z = self.u[2::pd]
        x_def = x_orig + scale_factor * disp_x; y_def = y_orig + scale_factor * disp_y; z_def = z_orig + scale_factor * disp_z
        fig = plt.figure(figsize=(8, 6)); ax = fig.add_subplot(111, projection='3d')
        plt.title(f"Static Analysis (Scale: {scale_factor})", fontsize=10); ax.set_xlabel("X", fontsize=8); ax.set_ylabel("Y", fontsize=8); ax.set_zlabel("Z", fontsize=8)
        for n1, n2 in self.conn:
            ax.plot([x_orig[n1], x_orig[n2]], [y_orig[n1], y_orig[n2]], [z_orig[n1], z_orig[n2]], 'k-', lw=1, alpha=0.3)
            ax.plot([x_def[n1], x_def[n2]], [y_def[n1], y_def[n2]], [z_def[n1], z_def[n2]], 'b--', lw=1.5)
        sc = ax.scatter(x_def, y_def, z_def, c=self.smoothed_stresses, cmap='jet', s=35, edgecolor='k')
        cbar = fig.colorbar(sc, ax=ax, shrink=0.7, pad=0.1); cbar.set_label('Stress (Pa)', rotation=270, labelpad=15)
        plt.tight_layout()
        return fig

    def _create_modal_plot_figure(self, mode_num):
        """모드 해석 결과를 Matplotlib Figure 객체로 생성하여 반환합니다."""
        if self.mode_shapes is None or self.mode_shapes.shape[1] < mode_num: return None
        
        mode_index = mode_num - 1
        mode_shape_vector = self.mode_shapes[:, mode_index]
        freq_hz = self.natural_frequencies[mode_index] / (2 * np.pi)
        pd = 6
        try: scale_factor = float(self.scale_factor_input.text())
        except (AttributeError, ValueError): scale_factor = 50.0
        
        x_orig, y_orig, z_orig = self.points[:,0], self.points[:,1], self.points[:,2]
        disp_x = mode_shape_vector[0::pd]; disp_y = mode_shape_vector[1::pd]; disp_z = mode_shape_vector[2::pd]
        disp_mag = np.sqrt(disp_x**2 + disp_y**2 + disp_z**2)
        max_mag = np.max(disp_mag)
        if max_mag > 0: disp_x, disp_y, disp_z = disp_x/max_mag, disp_y/max_mag, disp_z/max_mag
        x_def = x_orig + scale_factor * disp_x; y_def = y_orig + scale_factor * disp_y; z_def = z_orig + scale_factor * disp_z
        
        fig = plt.figure(figsize=(8, 6)); ax = fig.add_subplot(111, projection='3d')
        plt.title(f"Mode #{mode_num} ({freq_hz:.2f} Hz)", fontsize=10); ax.set_xlabel("X", fontsize=8); ax.set_ylabel("Y", fontsize=8); ax.set_zlabel("Z", fontsize=8)
        
        for n1, n2 in self.conn:
            ax.plot([x_orig[n1], x_orig[n2]], [y_orig[n1], y_orig[n2]], [z_orig[n1], z_orig[n2]], 'k-', lw=1, alpha=0.3)
            ax.plot([x_def[n1], x_def[n2]], [y_def[n1], y_def[n2]], [z_def[n1], z_def[n2]], 'r--', lw=2)
        sc = ax.scatter(x_def, y_def, z_def, c=disp_mag, cmap='viridis', s=35, edgecolor='k')
        cbar = fig.colorbar(sc, ax=ax, shrink=0.7, pad=0.1); cbar.set_label('Relative Displacement')
        plt.tight_layout()
        return fig

    def plot_static_results(self):
        if self.u is None: QMessageBox.warning(self, "No Data", "Please run the simulation first."); return
        fig = self._create_static_plot_figure()
        if fig: plt.show()
            
    def plot_modal_results(self):
        if self.mode_shapes is None: QMessageBox.warning(self, "No Data", "Please run the simulation first."); return
        num_modes = self.mode_shapes.shape[1]
        mode_num, ok = QInputDialog.getInt(self, "Select Mode Shape", f"Enter mode number (1 to {num_modes}):", 1, 1, num_modes)
        if ok:
            fig = self._create_modal_plot_figure(mode_num)
            if fig: plt.show()

    def get_timoshenko_stiffness_matrix(self, L, E, G, A, I_x, I_y, J, kappa_y, kappa_z):
        phi_z = (12 * E * I_y) / (G * kappa_y * A * L**2) if (G * kappa_y * A * L**2) > 0 else 0
        phi_y = (12 * E * I_x) / (G * kappa_z * A * L**2) if (G * kappa_z * A * L**2) > 0 else 0
        k11_z = (12*E*I_y)/(L**3*(1+phi_z)) if L > 0 else 0; k12_z = (6*E*I_y)/(L**2*(1+phi_z)) if L > 0 else 0
        k22_z = ((4+phi_z)*E*I_y)/(L*(1+phi_z)) if L > 0 else 0; k23_z = ((2-phi_z)*E*I_y)/(L*(1+phi_z)) if L > 0 else 0
        k11_y = (12*E*I_x)/(L**3*(1+phi_y)) if L > 0 else 0; k12_y = (6*E*I_x)/(L**2*(1+phi_y)) if L > 0 else 0
        k22_y = ((4+phi_y)*E*I_x)/(L*(1+phi_y)) if L > 0 else 0; k23_y = ((2-phi_y)*E*I_x)/(L*(1+phi_y)) if L > 0 else 0
        torsion_stiffness = G*J/L if L>0 else 0
        return np.array([
            [A*E/L if L>0 else 0,0,0,0,0,0,-A*E/L if L>0 else 0,0,0,0,0,0], [0,k11_z,0,0,0,k12_z,0,-k11_z,0,0,0,k12_z],
            [0,0,k11_y,0,-k12_y,0,0,0,-k11_y,0,-k12_y,0], [0,0,0,torsion_stiffness,0,0,0,0,0,-torsion_stiffness,0,0],
            [0,0,-k12_y,0,k22_y,0,0,0,k12_y,0,k23_y,0], [0,k12_z,0,0,0,k22_z,0,-k12_z,0,0,0,k23_z],
            [-A*E/L if L>0 else 0,0,0,0,0,0,A*E/L if L>0 else 0,0,0,0,0,0], [0,-k11_z,0,0,0,-k12_z,0,k11_z,0,0,0,-k12_z],
            [0,0,-k11_y,0,k12_y,0,0,0,k11_y,0,k12_y,0], [0,0,0,-torsion_stiffness,0,0,0,0,0,torsion_stiffness,0,0],
            [0,0,-k12_y,0,k23_y,0,0,0,k12_y,0,k22_y,0], [0,k12_z,0,0,0,k23_z,0,-k12_z,0,0,0,k22_z]])
    
    def get_lumped_mass_matrix(self, L, A, Ix, Iy, J, rho):
        """
        3D 빔 요소에 대한 집중 질량 행렬을 반환합니다.
        """
        m_el = np.zeros((12, 12))
        trans_mass = rho * A * L / 2
        rot_mass_x = rho * J * L / 2
        rot_mass_y = rho * Ix * L / 2
        rot_mass_z = rho * Iy * L / 2
        m_el[0, 0] = trans_mass; m_el[1, 1] = trans_mass; m_el[2, 2] = trans_mass
        m_el[3, 3] = rot_mass_x; m_el[4, 4] = rot_mass_y; m_el[5, 5] = rot_mass_z
        m_el[6, 6] = trans_mass; m_el[7, 7] = trans_mass; m_el[8, 8] = trans_mass
        m_el[9, 9] = rot_mass_x; m_el[10, 10] = rot_mass_y; m_el[11, 11] = rot_mass_z
        return m_el

    def bc_nodes_indexing(self, element_type, bc_name):
        try:
            cells = self.mesh.cells_dict.get(element_type)
            if cells is None: return np.array([], dtype=int)
            phys_data_map = self.mesh.cell_data_dict.get("gmsh:physical", {}); phys_data = phys_data_map.get(element_type)
            if phys_data is None: return np.array([], dtype=int)
            target_id = self.mesh.field_data[bc_name][0]
            return np.unique(cells[phys_data == target_id].flatten())
        except (KeyError, IndexError):
            QMessageBox.warning(self, "Warning", f"Physical group '{bc_name}' not found."); return np.array([], dtype=int)

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = BeamAnalysisWindow()
    window.show()
    sys.exit(app.exec_())

